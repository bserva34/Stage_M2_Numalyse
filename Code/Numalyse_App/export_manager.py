from PySide6.QtWidgets import QWidget, QVBoxLayout, QPushButton, QFileDialog, QRadioButton, QLabel, QLineEdit, QDialog, QButtonGroup, QHBoxLayout, QApplication
from PySide6.QtCore import Qt

import json
import os
import cv2
import numpy as np
import tempfile
import textwrap

from moviepy import VideoFileClip, AudioFileClip
from PIL import Image, ImageDraw, ImageFont
from io import BytesIO

# Pour PDF
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import Paragraph, SimpleDocTemplate
from reportlab.platypus import Image as PDFImage
# Pour DOCX
from docx import Document  
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
# Pour ODT
from odf.opendocument import OpenDocumentText  
from odf.text import P, LineBreak
from odf.style import Style, TextProperties, ParagraphProperties, GraphicProperties
from odf.draw import Frame
from odf.draw import Image as ODTImage


# Importation de classes
from message_popup import MessagePopUp
from time_manager import TimeManager
from exportvideo_thread import ExportVideoThread
from exporttext_thread import ExportTextThread
from no_focus_push_button import NoFocusPushButton



class ExportManager(QWidget):
    # Définition des styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'TitleStyle',
        parent=styles['Title'],
        fontSize=18,
        textColor=colors.red,
        spaceAfter=20
    )
    subtitle_style = ParagraphStyle(
        'SubtitleStyle',
        parent=styles['Heading2'],
        fontSize=12,
        textColor=colors.blue,
        spaceAfter=10
    )
    note_style = ParagraphStyle(
        'NoteStyle',
        parent=styles['BodyText'],
        fontSize=10,
        textColor=colors.black,
        leftIndent=20
    )

    def __init__(self, parent=None, vlc=None, projectmanager=None, format_export_text=[]):
        super().__init__(parent)
        self.seg = parent
        self.vlc = vlc
        self.project_manager = projectmanager

        self.file_path = self.project_manager.project_path
        self.title = self.project_manager.project_name
        self.time_manager = TimeManager(fps=self.vlc.fps)
        self.format_export_text = format_export_text

        self.configure()

    def configure(self):
        """ Ouvre une fenêtre de configuration pour choisir le mode. """
        dialog = QDialog(self)
        dialog.setWindowTitle("Exportation du Travail")
        dialog_layout = QVBoxLayout(dialog)

        # Options pour le type d'exportation
        num_label = QLabel("Type d'exportation :", dialog)
        dialog_layout.addWidget(num_label)

        num_group = QButtonGroup(dialog)
        option_1 = QRadioButton("Fichier texte", dialog)
        option_2 = QRadioButton("Vidéo annoté", dialog)
        num_group.addButton(option_1)
        num_group.addButton(option_2)
        option_1.setChecked(True)
        dialog_layout.addWidget(option_1)
        dialog_layout.addWidget(option_2)

        dialog_load = QHBoxLayout()
        load = QLabel("")
        load.setStyleSheet("color: blue;")
        load.setAlignment(Qt.AlignCenter)
        dialog_load.addWidget(load)
        dialog_layout.addLayout(dialog_load)

        # Boutons OK/Annuler
        button_layout = QHBoxLayout()
        ok_button = NoFocusPushButton("OK", dialog)
        cancel_button = NoFocusPushButton("Annuler", dialog)
        button_layout.addWidget(ok_button)
        button_layout.addWidget(cancel_button)
        dialog_layout.addLayout(button_layout)

        def on_ok():
            if self.file_path:
                load.setText("Exportation en cours ⌛")
                QApplication.processEvents()
                if option_1.isChecked():
                    # self.export_thread = ExportTextThread(self,option=4)
                    # self.export_thread.segmentation_done.connect(lambda: export_done(dialog))
                    # self.export_thread.start()
                    if self.format_export_text[0]:
                        self.export_thread = ExportTextThread(self,option=1)
                        self.export_thread.segmentation_done.connect(lambda: export_done(dialog))
                        self.export_thread.start()
                    elif self.format_export_text[1]:
                        self.export_thread = ExportTextThread(self,option=2)
                        self.export_thread.segmentation_done.connect(lambda: export_done(dialog))
                        self.export_thread.start()
                    elif self.format_export_text[2]:
                        self.export_thread = ExportTextThread(self,option=3)
                        self.export_thread.segmentation_done.connect(lambda: export_done(dialog))
                        self.export_thread.start()
                else:
                    # Lancement de l'export vidéo dans un thread séparé
                    self.export_thread = ExportVideoThread(self)
                    # Connexion du signal segmentation_done pour savoir quand l'export est terminé
                    self.export_thread.segmentation_done.connect(lambda: export_done(dialog))
                    self.export_thread.start()

        def on_cancel():
            # Si un export est en cours, demande l'annulation
            if hasattr(self, 'export_thread'):
                print('annulation')
                self.export_thread.stop()
            dialog.reject()

        def export_done(dialog):
            if(self.export_thread.running):
                affichage = MessagePopUp(self) 
                dialog.accept()

        ok_button.clicked.connect(on_ok)
        cancel_button.clicked.connect(on_cancel)

        dialog.setLayout(dialog_layout)
        dialog.exec()

    def get_images(self):
        stock_images = []
        stock_frames = [btn_data["frame1"] + 10 for btn_data in self.seg.display.stock_button]

        cap = cv2.VideoCapture(self.vlc.path_of_media)
        if not cap.isOpened():
            print("Impossible d'ouvrir la vidéo.")
            return

        frame_idx = 0
        stock_frames_set = set(stock_frames)

        while True:
            ret, frame = cap.read()
            if not ret:
                break
            if frame_idx in stock_frames_set:
                stock_images.append(frame)
            frame_idx += 1
        cap.release()

        return stock_images

    def size_of_img(self,img):
        max_width = 300
        max_height = 200

        _, img_bytes = cv2.imencode('.png', img)
        img_stream = BytesIO(img_bytes.tobytes())
        height, width, _ = img.shape
        width_ratio = max_width / float(width)
        height_ratio = max_height / float(height)
        ratio = min(width_ratio, height_ratio, 1.0)
        new_width = int(width * ratio)
        new_height = int(height * ratio)
        return img_stream,new_width,new_height

    def export_txt(self, callback=None):
        self.file_path = os.path.join(self.file_path, f"{self.title}.txt")  # Correction de l'extension

        try:
            with open(self.file_path, "w", encoding="utf-8") as fichier:
                total_plans = len(self.seg.display.stock_button)
                fichier.write(f"{total_plans}\n")

                for btn_data in self.seg.display.stock_button:
                    if not callback():
                        print("Exportation annulée par l'utilisateur.")
                        return
                    button = btn_data["button"]
                    line = f"{button.text()}"
                    fichier.write(line + "\n")

            print(f"Fichier TXT enregistré : {self.file_path}")
        except Exception as e:
            print(f"Erreur lors de l'exportation TXT : {e}")


    def export_pdf(self,callback=None):
        self.file_path = os.path.join(self.file_path, f"{self.title}.pdf")

        stock_images=self.get_images()

        try:
            doc = SimpleDocTemplate(self.file_path, pagesize=A4)
            elements = []
            elements.append(Paragraph("Étude cinématographique", self.title_style))

            for idx, btn_data in enumerate(self.seg.display.stock_button):
                if not callback():
                    print("Exportation annulée par l'utilisateur.")
                    return
                button = btn_data["button"]
                time_str = self.time_manager.m_to_hmsf(btn_data["time"])
                end_str = self.time_manager.m_to_hmsf(btn_data["end"] - btn_data["time"])
                elements.append(Paragraph(f"- {button.text()} -> Début : {time_str} / Durée : {end_str}", self.subtitle_style))
                for note_widget in self.seg.display.button_notes.get(button, []):
                    note_text = note_widget.toPlainText()
                    self.put_multiline_text(elements, note_text)
                if idx < len(stock_images):
                    img = stock_images[idx]
                    img_stream,new_width,new_height=self.size_of_img(img)
                    img_obj = PDFImage(img_stream, width=new_width, height=new_height)
                    elements.append(img_obj)
            doc.build(elements)
            print(f"Fichier PDF enregistré : {self.file_path}")
        except Exception as e:
            print(f"Erreur lors de l'exportation PDF : {e}")

    def export_docx(self, callback=None):
        self.file_path = os.path.join(self.file_path, f"{self.title}.docx")
        try:
            doc = Document()
            # Ajout du titre centré
            titre_paragraphe = doc.add_paragraph()
            titre_paragraphe.alignment = WD_ALIGN_PARAGRAPH.CENTER
            titre_run = titre_paragraphe.add_run("Étude cinématographique")
            titre_run.font.size = Pt(24)
            titre_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            
            # Récupérer les images comme dans export_pdf
            stock_images = self.get_images()
        
            # Boucle sur chaque bouton et note associée
            for idx, btn_data in enumerate(self.seg.display.stock_button):
                if not callback():
                    print("Exportation annulée par l'utilisateur.")
                    return
                button = btn_data["button"]
                time_str = self.time_manager.m_to_hmsf(btn_data["time"])
                end_str = self.time_manager.m_to_hmsf(btn_data["end"] - btn_data["time"])
                doc.add_heading(f"- {button.text()} -> Début : {time_str} / Durée : {end_str}", level=2)
                
                for note_widget in self.seg.display.button_notes.get(button, []):
                    note_text = note_widget.toPlainText()
                    doc.add_paragraph(note_text)
                    
                # Ajout de l'image associée (si disponible)
                if idx < len(stock_images):
                    img = stock_images[idx]
                    img_stream,new_width,new_height=self.size_of_img(img)
                    img_docx_width = Inches(new_width / 96.0)
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.add_run().add_picture(img_stream, width=img_docx_width)
            
            doc.save(self.file_path)
            print(f"Fichier DOCX enregistré : {self.file_path}")
        except Exception as e:
            print(f"Erreur lors de l'exportation DOCX : {e}")

    def export_odt(self, callback=None):
        self.file_path = os.path.join(self.file_path, f"{self.title}.odt")
        try:
            doc = OpenDocumentText()

            # Création d'un style pour le titre
            title_style = Style(name="Titre", family="paragraph")
            title_style.addElement(ParagraphProperties(textalign="center"))
            title_style.addElement(TextProperties(color="#FF0000", fontsize="24pt"))
            doc.styles.addElement(title_style)

            # Ajout du titre avec le style défini
            titre_paragraphe = P(stylename=title_style, text="Étude cinématographique")
            doc.text.addElement(titre_paragraphe)


            first_button_style = Style(name="PremierBouton", family="paragraph")
            first_button_style.addElement(TextProperties(color="#0000FF", fontsize="14pt"))
            doc.styles.addElement(first_button_style)

            # Parcours des boutons et ajout du contenu
            for idx, btn_data in enumerate(self.seg.display.stock_button):
                if not callback():
                    print("Exportation annulée par l'utilisateur.")
                    return
                button = btn_data["button"]
                time_str = self.time_manager.m_to_hmsf(btn_data["time"])
                end_str = self.time_manager.m_to_hmsf(btn_data["end"] - btn_data["time"])
                doc.text.addElement(P(stylename=first_button_style,text=f"- {button.text()} -> Début : {time_str} / Durée : {end_str}"))
                for note_widget in self.seg.display.button_notes.get(button, []):
                    note_text = note_widget.toPlainText()
                    doc.text.addElement(P(text=note_text))
                doc.text.addElement(P())
            doc.save(self.file_path)
            print(f"Fichier ODT enregistré : {self.file_path}")
        except Exception as e:
            print(f"Erreur lors de l'exportation ODT : {e}")


    def put_multiline_text(self, elements, text):
        lines = text.split("\n")
        for line in lines:
            line = line.replace("\t", "&nbsp;&nbsp;&nbsp;&nbsp;")
            elements.append(Paragraph(line, self.note_style))

    def export_video(self, callback=None):
        try:
            self.file_path = os.path.join(self.file_path, f"{self.title}.mp4")
            self.project_manager.path_of_super = self.file_path

            temp_dir = tempfile.gettempdir()
            temp_video_path = os.path.join(temp_dir, "temp_video.mp4")

            cap = cv2.VideoCapture(self.vlc.path_of_media)
            if not cap.isOpened():
                print("Impossible d'ouvrir la vidéo.")
                return

            fourcc = cv2.VideoWriter_fourcc(*"mp4v")
            out = cv2.VideoWriter(temp_video_path, fourcc, 30.0, (int(cap.get(3)), int(cap.get(4))))

            cpt = 0
            while cap.isOpened():
                if not callback():
                    print("Exportation annulée par l'utilisateur.")
                    cap.release()
                    out.release()
                    if os.path.exists(temp_video_path):
                        os.remove(temp_video_path)
                    return

                ret, frame = cap.read()
                if not ret:
                    break

                active_texts = [
                    btn_data for btn_data in self.seg.display.stock_button
                    if btn_data["frame1"] <= cpt < btn_data["frame2"]
                ]
                for btn_data in active_texts:
                    button = btn_data["button"]
                    time_str = self.time_manager.m_to_hmsf(btn_data["time"])
                    end_str = self.time_manager.m_to_hmsf(btn_data["end"] - btn_data["time"])
                    txt = button.text()
                    txt2 = f"Debut : {time_str} / Duree : {end_str}"
                    #txt3 = [note_widget.toPlainText() for note_widget in self.seg.display.button_notes.get(button, [])]
                    height, width, _ = frame.shape
                    self.write_text_horizontal_on_video(frame, txt, txt2,"", width)
                out.write(frame)
                cpt += 1

            cap.release()
            out.release()

            if not callback():
                print("Exportation annulée par l'utilisateur.")
                if os.path.exists(temp_video_path):
                    os.remove(temp_video_path)
                return

            video_clip = VideoFileClip(temp_video_path)
            try:
                audio_clip = AudioFileClip(self.vlc.path_of_media)
                final_clip = video_clip.set_audio(audio_clip)
            except Exception as e:
                print(f"⚠️ Impossible de charger l'audio : {e}")
                final_clip = video_clip

            final_clip.write_videofile(self.file_path, codec="libx264", audio_codec="aac",logger=None)
            video_clip.close()
            audio_clip.close()
            final_clip.close()

            os.remove(temp_video_path)
        except Exception as e:
            print(f"Erreur pendant l'export vidéo : {e}")

    def write_text_horizontal_on_video(self, frame, txt, txt2, txt3, max_width, line_spacing=24):
        final_txt = txt + " - " + txt2
        for text in txt3:
            lines = text.split("\n")
            final_txt += " - " + " / ".join(line.replace("\t", "   ") for line in lines)

        font_scale = 0.7
        thickness_outline = 4
        thickness_text = 1
        font = cv2.FONT_HERSHEY_SIMPLEX

        char_size = cv2.getTextSize("A", font, font_scale, thickness_text)[0][0]
        max_chars_per_line = (max_width + 50) // char_size

        wrapped_text = textwrap.fill(final_txt, width=max_chars_per_line)
        x, y = 50, 30
        for line in wrapped_text.split("\n"):
            cv2.putText(frame, line, (x, y), font, font_scale, (0, 0, 0), thickness_outline, cv2.LINE_AA)
            cv2.putText(frame, line, (x, y), font, font_scale, (255, 255, 255), thickness_text, cv2.LINE_AA)
            y += line_spacing

    def write_text_on_video(self, frame, txt, txt2, txt3, decalage):
        cv2.putText(frame, txt, (50, decalage + 50), cv2.FONT_HERSHEY_SIMPLEX, 1, (0, 0, 0), 5, cv2.LINE_AA)
        cv2.putText(frame, txt, (50, decalage + 50), cv2.FONT_HERSHEY_SIMPLEX, 1, (255, 255, 255), 2, cv2.LINE_AA)
        cv2.putText(frame, txt2, (50, decalage + 80), cv2.FONT_HERSHEY_SIMPLEX, 0.7, (0, 0, 0), 5, cv2.LINE_AA)
        cv2.putText(frame, txt2, (50, decalage + 80), cv2.FONT_HERSHEY_SIMPLEX, 0.7, (255, 255, 255), 2, cv2.LINE_AA)
        val = decalage + 80
        for i, text in enumerate(txt3):
            val = decalage + 110 + (i * 20)
            self.put_multiline_text_video(frame, text, (50, val))
        return val + 10

    def put_multiline_text_video(self, frame, text, position, font_scale=0.5, line_spacing=15):
        lines = text.split("\n")
        x, y = position
        for i, line in enumerate(lines):
            cv2.putText(frame, line.replace("\t", "    "), (x, y + i * line_spacing),
                        cv2.FONT_HERSHEY_SIMPLEX, font_scale, (0, 0, 0), 5, cv2.LINE_AA)
            cv2.putText(frame, line.replace("\t", "    "), (x, y + i * line_spacing),
                        cv2.FONT_HERSHEY_SIMPLEX, font_scale, (255, 255, 255), 2, cv2.LINE_AA)
